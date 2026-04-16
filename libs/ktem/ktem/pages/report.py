import json
import os
import re
import shutil
import tempfile
from datetime import datetime
from pathlib import Path

import gradio as gr
from ktem.app import BasePage


def parse_reference_doc(file_path):
    """解析参考报告，提取章节结构和完整内容"""
    try:
        from docx import Document
        doc = Document(file_path)
        chapters = []
        current = None

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            style = para.style.name
            if style.startswith("Heading") or style.startswith("标题"):
                try:
                    level = int(re.search(r'\d', style).group())
                except Exception:
                    level = 1
                current = {
                    "title": para.text.strip(),
                    "level": level,
                    "content": "",
                    "style": style,
                }
                chapters.append(current)
            elif current is not None:
                current["content"] += para.text.strip() + "\n"

        for ch in chapters:
            ch["content_preview"] = ch["content"][:200]

        return chapters, None
    except Exception as e:
        return None, f"参考报告解析失败：{e}"


def parse_data_files(data_files):
    """解析上传的数据文件，返回结构化摘要"""
    if not data_files:
        return []

    files = data_files if isinstance(data_files, list) else [data_files]
    results = []

    for f in files:
        file_path = f.name if hasattr(f, "name") else f
        fname = Path(file_path).name
        ext = Path(file_path).suffix.lower()

        try:
            if ext == ".csv":
                import csv
                with open(file_path, "r", encoding="utf-8", errors="ignore") as fp:
                    reader = csv.reader(fp)
                    rows = list(reader)
                headers = rows[0] if rows else []
                preview_rows = rows[1:6] if len(rows) > 1 else []
                summary = f"文件名：{fname}\n列名：{', '.join(headers)}\n数据行数：{len(rows)-1}\n前5行：\n"
                for row in preview_rows:
                    summary += "  " + ", ".join(str(x) for x in row) + "\n"
                results.append({"name": fname, "type": "csv", "summary": summary})

            elif ext in [".xlsx", ".xls"]:
                try:
                    import pandas as pd
                    xl = pd.ExcelFile(file_path)
                    summary = f"文件名：{fname}\n工作表：{', '.join(xl.sheet_names)}\n"
                    for sheet in xl.sheet_names:
                        df = pd.read_excel(file_path, sheet_name=sheet)
                        summary += f"\n【{sheet}】列名：{', '.join(df.columns.astype(str))}\n"
                        summary += f"行数：{len(df)}\n前3行：\n{df.head(3).to_string(index=False)}\n"
                    results.append({"name": fname, "type": "excel", "summary": summary})
                except ImportError:
                    results.append({"name": fname, "type": "excel",
                                    "summary": f"文件名：{fname}\n（需安装pandas）"})

            elif ext == ".json":
                with open(file_path, "r", encoding="utf-8", errors="ignore") as fp:
                    data = json.load(fp)
                summary = f"文件名：{fname}\n"
                if isinstance(data, list):
                    summary += f"条数：{len(data)}\n"
                    if data:
                        summary += f"字段：{', '.join(str(k) for k in data[0].keys())}\n"
                    summary += f"前2条：{json.dumps(data[:2], ensure_ascii=False)[:300]}\n"
                else:
                    summary += f"内容：{json.dumps(data, ensure_ascii=False)[:500]}\n"
                results.append({"name": fname, "type": "json", "summary": summary})

            else:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as fp:
                    text = fp.read()[:800]
                results.append({"name": fname, "type": "text",
                                "summary": f"文件名：{fname}\n内容：{text}"})
        except Exception as e:
            results.append({"name": fname, "type": "unknown",
                            "summary": f"文件名：{fname}\n读取失败：{e}"})

    return results


def parse_exp_desc(file_obj):
    """解析本次实验说明文件（Word或TXT）"""
    if file_obj is None:
        return ""
    try:
        file_path = file_obj.name if hasattr(file_obj, "name") else file_obj
        ext = Path(file_path).suffix.lower()
        if ext == ".docx":
            from docx import Document
            doc = Document(file_path)
            text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        else:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        return text[:3000]  # 限制长度
    except Exception as e:
        return f"（实验说明读取失败：{e}）"


def generate_report_content(chapters, data_summaries, image_count,
                             exp_name, exp_date, exp_person, exp_purpose,
                             exp_desc=""):
    """调用大模型生成报告内容"""
    try:
        from kotaemon.base import HumanMessage
        from ktem.llms.manager import llms
        llm = llms.get_default()
        if llm is None:
            return None, "未配置默认大模型，请在资源配置中设置"
    except Exception as e:
        return None, f"大模型初始化失败：{e}"

    # 参考报告结构（限制长度避免 token 超限）
    ref_structure = ""
    for ch in chapters:
        indent = "  " * (ch["level"] - 1)
        ref_structure += f"{indent}{'#'*ch['level']} {ch['title']}\n"
        preview = ch.get("content_preview", "")[:150]
        if preview:
            ref_structure += f"{indent}原文参考：{preview}\n\n"

    # 数据摘要（限制总长度）
    data_desc = ""
    for i, d in enumerate(data_summaries):
        chunk = d["summary"][:300]
        data_desc += f"\n数据{i+1}：\n{chunk}\n"

    img_desc = f"共 {image_count} 张图表，按顺序插入合适位置" if image_count > 0 else "无图表"

    exp_desc_section = f"""
本次实验详细说明（优先参考此内容改写各章节，不仅限于数据章节）：
{exp_desc}
""" if exp_desc else ""

    prompt = f"""你是仿真实验报告撰写工程师。请为本次实验生成报告。

参考报告结构（请保留章节结构，仅修改与实验相关的具体内容）：
{ref_structure}

本次实验信息：
- 实验名称：{exp_name}
- 实验日期：{exp_date}
- 实验人员：{exp_person}（必须原样使用，禁止修改）
- 实验目的：{exp_purpose}
{exp_desc_section}
实验数据：
{data_desc if data_desc else "无数据文件"}

图表：{img_desc}

要求：
1. 保留参考报告的所有章节，标题中的旧实验名称改为本次实验名称
2. 优先根据【本次实验详细说明】改写对应章节（包括实验背景、场景、参数设置等），不仅限于数据章节
3. 有实验数据的章节结合数据分析改写；既无说明也无数据的章节直接复用参考报告原文
4. 实验人员"{exp_person}"必须原样使用，绝对禁止替换
5. 统一使用"实验"而非"试验"
6. 每章content控制在300字以内，简洁专业

注意：content字段内如需表达子标题，用##或###标记，表格用Markdown格式（|列1|列2|），列表用-开头。

输出纯JSON，无任何其他内容：
{{"title":"标题","chapters":[{{"level":1,"title":"章节标题","content":"内容","has_image":false,"image_index":null}}]}}"""

    raw = ""
    try:
        print(f"[报告生成] prompt字符数：{len(prompt)}，估算token：{len(prompt)//2}")
        response = llm([HumanMessage(content=prompt)])
        raw = response.text.strip()
        raw = re.sub(r'^```json\s*', '', raw)
        raw = re.sub(r'```\s*$', '', raw)
        raw = raw.strip()
        result = json.loads(raw)
        return result, None
    except json.JSONDecodeError as e:
        preview = raw[:300] if raw else "（无输出）"
        return None, f"JSON解析失败：{e}\n输出片段：{preview}"
    except Exception as e:
        preview = raw[:300] if raw else "（无输出）"
        return None, f"大模型调用失败：{e}\n输出片段：{preview}"


def build_word_doc(reference_path, report_content, image_paths, output_path,
                   exp_name="", exp_date="", exp_person=""):
    """基于参考报告样式生成新 Word 文档"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import copy

    # 复制参考报告（继承所有样式定义）
    shutil.copy(reference_path, output_path)
    doc = Document(output_path)

    # 清空正文内容，保留样式
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("p", "tbl", "sectPr"):
            if tag != "sectPr":  # 保留节属性（页面设置）
                body.remove(child)

    def add_para(text, style_name="Normal", align=None, bold=False, size=None):
        """添加段落，使用参考报告样式"""
        p = doc.add_paragraph(style=style_name)
        run = p.add_run(text)
        if bold:
            run.font.bold = True
        if size:
            run.font.size = Pt(size)
        if align:
            p.alignment = align
        # 首行缩进（正文段落）
        if style_name == "Normal":
            p.paragraph_format.first_line_indent = Pt(22)
        return p

    def add_heading_para(text, level):
        """添加标题，使用Word内置标题样式"""
        try:
            style_name = f"Heading {level}"
            p = doc.add_paragraph(style=style_name)
            p.add_run(text)
        except Exception:
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(14 - level)
        return p

    def add_info_table(data):
        """添加封面信息表格"""
        t = doc.add_table(rows=len(data), cols=2)
        try:
            t.style = "Table Grid"
        except Exception:
            pass
        for i, (label, value) in enumerate(data):
            t.rows[i].cells[0].text = label
            t.rows[i].cells[1].text = str(value)
            for cell in t.rows[i].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(11)
        return t

    # ── 封面 ──
    title_p = doc.add_paragraph(style="Normal")
    title_run = title_p.add_run(report_content.get("title", "仿真实验报告"))
    title_run.font.bold = True
    title_run.font.size = Pt(16)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.first_line_indent = Pt(0)
    doc.add_paragraph("")

    add_info_table([
        ("实验名称", exp_name),
        ("实验日期", exp_date),
        ("实验人员", exp_person),
        ("报告生成时间", datetime.now().strftime("%Y年%m月%d日 %H:%M")),
    ])
    doc.add_paragraph("")

    # ── 各章节 ──
    for ch in report_content.get("chapters", []):
        level = ch.get("level", 1)
        title = ch.get("title", "")
        content = ch.get("content", "")

        add_heading_para(title, level)

        if content:
            lines = content.split("\n")
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                if not line:
                    i += 1
                    continue

                # 检测 Markdown 表格（以|开头）
                if line.startswith("|"):
                    # 收集连续的表格行
                    table_lines = []
                    while i < len(lines) and lines[i].strip().startswith("|"):
                        table_lines.append(lines[i].strip())
                        i += 1
                    # 过滤分隔行（|---|）
                    data_lines = [l for l in table_lines
                                  if not all(c in "|-: " for c in l)]
                    if data_lines:
                        # 解析表格
                        rows = []
                        for tl in data_lines:
                            cells = [c.strip() for c in tl.strip("|").split("|")]
                            rows.append(cells)
                        if rows:
                            ncols = max(len(r) for r in rows)
                            t = doc.add_table(rows=len(rows), cols=ncols)
                            try:
                                t.style = "Table Grid"
                            except Exception:
                                pass
                            for ri, row in enumerate(rows):
                                for ci in range(ncols):
                                    val = row[ci] if ci < len(row) else ""
                                    cell = t.rows[ri].cells[ci]
                                    cell.text = val
                                    for para in cell.paragraphs:
                                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        for run in para.runs:
                                            run.font.size = Pt(10)
                                            if ri == 0:
                                                run.font.bold = True
                            doc.add_paragraph("")
                    continue

                # Markdown 标题（## 或 ###）
                if line.startswith("#"):
                    hashes = len(line) - len(line.lstrip("#"))
                    heading_text = line.lstrip("#").strip()
                    add_heading_para(heading_text, min(hashes, 4))
                    i += 1
                    continue

                # 列表项
                if line.startswith("- ") or line.startswith("• "):
                    p = doc.add_paragraph(line[2:], style="List Bullet")
                    for run in p.runs:
                        run.font.size = Pt(11)
                elif len(line) > 2 and line[0].isdigit() and line[1] in ".、":
                    p = doc.add_paragraph(line[2:].strip(), style="List Number")
                    for run in p.runs:
                        run.font.size = Pt(11)
                else:
                    p = doc.add_paragraph(style="Normal")
                    run = p.add_run(line)
                    run.font.size = Pt(11)
                    p.paragraph_format.first_line_indent = Pt(22)
                i += 1

        # 插入图片
        if ch.get("has_image") and ch.get("image_index"):
            idx = ch["image_index"] - 1
            if 0 <= idx < len(image_paths):
                try:
                    doc.add_picture(image_paths[idx], width=Inches(5.5))
                    cap = doc.add_paragraph(f"图{ch['image_index']}  {title}相关图表")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.paragraph_format.first_line_indent = Pt(0)
                    for run in cap.runs:
                        run.font.size = Pt(10)
                except Exception as e:
                    doc.add_paragraph(f"（图片插入失败：{e}）")

        doc.add_paragraph("")

    doc.save(output_path)
    return True, None


def generate_report(ref_file, exp_name, exp_date, exp_person,
                    exp_desc_file, data_paths, image_paths):
    exp_purpose = ""
    """报告生成主流程"""
    if ref_file is None:
        raise gr.Error("请上传参考报告（Word文档）")
    if not exp_name or not exp_name.strip():
        raise gr.Error("请填写实验名称")

    ref_path = ref_file.name if hasattr(ref_file, "name") else ref_file
    chapters, err = parse_reference_doc(ref_path)
    if err:
        raise gr.Error(err)
    if not chapters:
        raise gr.Error("参考报告中未找到章节结构，请确认文档使用了标题样式")

    # 从路径列表构造文件对象
    class FakePath:
        def __init__(self, p): self.name = p

    data_files = [FakePath(p) for p in (data_paths or [])]
    image_files = [FakePath(p) for p in (image_paths or [])]

    data_summaries = parse_data_files(data_files)
    exp_desc = parse_exp_desc(exp_desc_file)
    report_content, err = generate_report_content(
        chapters, data_summaries, len(image_files),
        exp_name.strip(),
        exp_date or datetime.now().strftime("%Y年%m月%d日"),
        exp_person.strip() if exp_person and exp_person.strip() else "未填写",
        "（请根据实验说明和数据推断，若无参考依据则写【未找到参考依据】）",
        exp_desc=exp_desc
    )
    if err:
        raise gr.Error(err)

    safe_name = "".join(c for c in exp_name if c.isalnum() or "\u4e00" <= c <= "\u9fff" or c in "._- ").strip()
    filename = f"{safe_name}_实验报告_{datetime.now().strftime('%Y%m%d%H%M')}.docx"
    # 使用 GRADIO_TEMP_DIR 确保文件可被下载
    gradio_tmp = os.environ.get("GRADIO_TEMP_DIR", tempfile.gettempdir())
    output_path = os.path.join(gradio_tmp, filename)

    ok, err = build_word_doc(
        ref_path, report_content,
        [p for p in (image_paths or [])],
        output_path,
        exp_name=exp_name.strip(),
        exp_date=exp_date or datetime.now().strftime("%Y年%m月%d日"),
        exp_person=exp_person.strip() if exp_person and exp_person.strip() else "未填写"
    )
    if err:
        raise gr.Error(f"文档生成失败：{err}")

    # 生成HTML下载链接
    fname = Path(output_path).name
    file_url = f"/file={output_path}"
    html = f'''<div style="padding:16px;background:#f0fdf4;border:1px solid #86efac;border-radius:8px;">
    <p style="margin:0 0 8px 0;color:#166534;font-weight:bold;">✅ 报告生成成功</p>
    <a href="{file_url}" download="{fname}"
       style="display:inline-flex;align-items:center;gap:6px;padding:8px 16px;
              background:#16a34a;color:white;border-radius:6px;text-decoration:none;font-size:14px;">
       📥 点击下载：{fname}
    </a>
</div>'''
    return html


class ReportPage(BasePage):
    """报告生成页面"""

    def __init__(self, app):
        self._app = app
        self.on_building_ui()

    def on_building_ui(self):
        gr.Markdown("# 📝 仿真实验报告生成")
        gr.Markdown(
            "上传参考报告（已有内容的 Word 文档），"
            "系统将学习其结构和风格，结合本次实验数据自动生成新报告。"
        )

        with gr.Row():
            # ── 左列：输入区
            with gr.Column(scale=1):
                gr.Markdown("### 📄 参考报告")
                self.ref_file = gr.File(
                    label="上传参考报告（.docx）",
                    file_types=[".docx"],
                )

                gr.Markdown("### 📋 本次实验信息")
                self.exp_name = gr.Textbox(
                    label="实验名称 *",
                    placeholder="请输入实验名称",
                )
                self.exp_date = gr.Textbox(
                    label="实验日期",
                    value=datetime.now().strftime("%Y年%m月%d日"),
                )
                self.exp_person = gr.Textbox(
                    label="实验人员",
                    placeholder="请输入实验人员姓名",
                )
                gr.Markdown("### 📝 本次实验说明（选填）")
                self.exp_desc_file = gr.File(
                    label="上传本次实验说明（.docx / .txt），系统优先参考此内容改写报告各章节",
                    file_types=[".docx", ".txt"],
                )

                gr.Markdown("### 📁 实验数据与图表")

                # 数据文件区
                self.data_files_state = gr.State([])
                with gr.Row():
                    self.data_upload = gr.File(
                        label="上传实验数据（CSV / Excel / JSON）",
                        file_count="multiple",
                        file_types=[".csv", ".xlsx", ".xls", ".json", ".txt"],
                        scale=3,
                    )
                    with gr.Column(scale=1, min_width=80):
                        self.add_data_btn = gr.Button("➕ 继续上传", size="sm")
                        self.clear_data_btn = gr.Button("🗑️ 清除全部", size="sm")
                self.data_list_check = gr.CheckboxGroup(
                    label="已加入数据文件（勾选后点「删除勾选」可单独删除）",
                    choices=[],
                    value=[],
                )
                self.del_data_btn = gr.Button("🗑️ 删除勾选的数据文件", size="sm")

                # 图表文件区
                self.image_files_state = gr.State([])
                with gr.Row():
                    self.image_upload = gr.File(
                        label="上传实验图表（PNG / JPG，按顺序插入报告）",
                        file_count="multiple",
                        file_types=[".png", ".jpg", ".jpeg"],
                        scale=3,
                    )
                    with gr.Column(scale=1, min_width=80):
                        self.add_img_btn = gr.Button("➕ 继续上传", size="sm")
                        self.clear_img_btn = gr.Button("🗑️ 清除全部", size="sm")
                self.image_list_check = gr.CheckboxGroup(
                    label="已加入图表文件（勾选后点「删除勾选」可单独删除，顺序即插入顺序）",
                    choices=[],
                    value=[],
                )
                self.del_img_btn = gr.Button("🗑️ 删除勾选的图表文件", size="sm")

                self.generate_btn = gr.Button(
                    "🚀 生成报告", variant="primary", size="lg"
                )

            # ── 右列：输出区
            with gr.Column(scale=1):
                gr.Markdown("### 📊 生成结果")
                self.output_file = gr.HTML(
                    value="<div style='padding:16px;color:#6b7280;'>报告生成后将在此处显示下载链接</div>",
                )
                gr.Markdown("""
### 💡 使用说明

**参考报告：**
- 格式为 Word（.docx），使用标准标题样式
- 推荐使用已有完整报告，系统最大程度复用其内容

**本次实验说明（选填）：**
- 格式为 Word（.docx）或 TXT
- 描述本次实验的背景、场景、参数配置等内容
- 系统优先参考此内容改写报告各章节，不仅限于数据章节
- 未上传时，无对应参考的章节将直接复用参考报告原文

**数据文件：**
- 支持 CSV、Excel（多工作表）、JSON
- 可多次点击「➕ 继续上传」追加文件
- 「🗑️ 清除全部」清空已加入列表

**图表文件：**
- 按顺序上传，依次插入报告对应位置
- 可多次点击「➕ 继续上传」追加，支持分批上传
- 已上传图表可勾选后点「删除勾选」单独删除

**生成逻辑：**
- 有实验说明的章节：优先参考实验说明改写
- 有实验数据的章节：结合数据分析改写
- 无对应内容的章节：直接复用参考报告原文
- 实验名称、日期、人员：严格使用填写内容，不做修改

**注意：**
- 请在「资源配置」中确认已设置默认大模型
- 生成约需 30～120 秒，请耐心等待
                """)

    def on_register_events(self):

        def append_files(new_files, state):
            if not new_files:
                choices = _choices(state)
                return state or [], gr.update(choices=choices, value=[])
            new_paths = [f.name if hasattr(f, "name") else f for f in new_files]
            existing = state or []
            merged = existing + [p for p in new_paths if p not in existing]
            choices = _choices(merged)
            return merged, gr.update(choices=choices, value=[])

        def clear_all_files():
            return [], None, gr.update(choices=[], value=[]), gr.update(choices=[], value=[])

        def del_checked_data(checked, state):
            """删除勾选的数据文件"""
            remaining = [p for p in (state or []) if Path(p).name not in checked]
            choices = _choices(remaining)
            return remaining, gr.update(choices=choices, value=[]), []

        def del_checked_img(checked, state):
            """删除勾选的图表文件"""
            remaining = [p for p in (state or []) if Path(p).name not in checked]
            choices = _choices(remaining)
            return remaining, gr.update(choices=choices, value=[]), []

        # 数据文件上传追加
        self.data_upload.change(
            fn=append_files,
            inputs=[self.data_upload, self.data_files_state],
            outputs=[self.data_files_state, self.data_list_check],
        )
        self.add_data_btn.click(
            fn=lambda: None,
            inputs=[],
            outputs=[self.data_upload],
        )
        self.clear_data_btn.click(
            fn=clear_all_files,
            inputs=[],
            outputs=[self.data_files_state, self.data_upload,
                     self.data_list_check, self.data_list_check],
        )
        self.del_data_btn.click(
            fn=del_checked_data,
            inputs=[self.data_list_check, self.data_files_state],
            outputs=[self.data_files_state, self.data_list_check, self.data_list_check],
        )

        # 图表文件上传追加
        self.image_upload.change(
            fn=append_files,
            inputs=[self.image_upload, self.image_files_state],
            outputs=[self.image_files_state, self.image_list_check],
        )
        self.add_img_btn.click(
            fn=lambda: None,
            inputs=[],
            outputs=[self.image_upload],
        )
        self.clear_img_btn.click(
            fn=clear_all_files,
            inputs=[],
            outputs=[self.image_files_state, self.image_upload,
                     self.image_list_check, self.image_list_check],
        )
        self.del_img_btn.click(
            fn=del_checked_img,
            inputs=[self.image_list_check, self.image_files_state],
            outputs=[self.image_files_state, self.image_list_check, self.image_list_check],
        )

        # 生成报告
        self.generate_btn.click(
            fn=generate_report,
            inputs=[
                self.ref_file,
                self.exp_name,
                self.exp_date,
                self.exp_person,
                self.exp_desc_file,
                self.data_files_state,
                self.image_files_state,
            ],
            outputs=[self.output_file],
        )


def _choices(paths):
    """返回CheckboxGroup的choices列表（显示文件名）"""
    if not paths:
        return []
    return [Path(p).name for p in paths]
